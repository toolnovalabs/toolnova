from pathlib import Path
import io
import json
import base64
import shutil
import uuid
import re
import secrets
import string
import zipfile
from urllib.parse import urlparse, parse_qs
from datetime import date, datetime

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, Response, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from pypdf import PdfReader, PdfWriter
from PIL import Image
from pillow_heif import register_heif_opener
from docx import Document
import openpyxl
import pandas as pd
import pdfplumber
import fitz
import qrcode
import markdown as md
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from app.seo_units import SEO_UNIT_PAGES

register_heif_opener()

# --------------------------------------------------
# PATHS
# --------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent

TEMPLATES_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"
TEMP_DIR = BASE_DIR / "temp"

SITEMAP_PATH = STATIC_DIR / "sitemap.xml"
ROBOTS_PATH = STATIC_DIR / "robots.txt"

FONTS_DIR = STATIC_DIR / "fonts"
DEJAVU_FONT_PATH = FONTS_DIR / "DejaVuSans.ttf"

STATIC_DIR.mkdir(exist_ok=True)
TEMP_DIR.mkdir(exist_ok=True)
FONTS_DIR.mkdir(exist_ok=True)

if DEJAVU_FONT_PATH.exists():
    pdfmetrics.registerFont(TTFont("DejaVuSans", str(DEJAVU_FONT_PATH)))

app = FastAPI(title="ToolNova")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

# --------------------------------------------------
# PLAN / PRICING LIMITS
# --------------------------------------------------
PLAN_LIMITS = {
    "free": {
        "merge_pdf_max_files": 10,
        "merge_pdf_max_total_mb": 60,
        "merge_pdf_max_total_bytes": 60 * 1024 * 1024,
        "merge_pdf_max_total_pages": 150,

        "pdf_general_max_mb": 40,
        "pdf_general_max_bytes": 40 * 1024 * 1024,
        "pdf_general_max_pages": 200,

        "pdf_convert_max_mb": 25,
        "pdf_convert_max_bytes": 25 * 1024 * 1024,
        "pdf_convert_max_pages": 100,

        "jpg_to_pdf_max_files": 10,
        "jpg_to_pdf_max_total_mb": 40,
        "jpg_to_pdf_max_total_bytes": 40 * 1024 * 1024,

        "heic_single_max_mb": 20,
        "heic_single_max_bytes": 20 * 1024 * 1024,
        "heic_bulk_max_files": 5,
        "heic_bulk_max_total_mb": 25,
        "heic_bulk_max_total_bytes": 25 * 1024 * 1024,

        "image_general_max_mb": 20,
        "image_general_max_bytes": 20 * 1024 * 1024,

        "office_general_max_mb": 15,
        "office_general_max_bytes": 15 * 1024 * 1024,

        "text_preview_max_chars": 20000,
        "csv_preview_max_rows": 100,
    },
    "pro": {
        "merge_pdf_max_files": 30,
        "merge_pdf_max_total_mb": 200,
        "merge_pdf_max_total_bytes": 200 * 1024 * 1024,
        "merge_pdf_max_total_pages": 1000,

        "pdf_general_max_mb": 150,
        "pdf_general_max_bytes": 150 * 1024 * 1024,
        "pdf_general_max_pages": 1000,

        "pdf_convert_max_mb": 100,
        "pdf_convert_max_bytes": 100 * 1024 * 1024,
        "pdf_convert_max_pages": 500,

        "jpg_to_pdf_max_files": 50,
        "jpg_to_pdf_max_total_mb": 200,
        "jpg_to_pdf_max_total_bytes": 200 * 1024 * 1024,

        "heic_single_max_mb": 50,
        "heic_single_max_bytes": 50 * 1024 * 1024,
        "heic_bulk_max_files": 50,
        "heic_bulk_max_total_mb": 250,
        "heic_bulk_max_total_bytes": 250 * 1024 * 1024,

        "image_general_max_mb": 75,
        "image_general_max_bytes": 75 * 1024 * 1024,

        "office_general_max_mb": 100,
        "office_general_max_bytes": 100 * 1024 * 1024,

        "text_preview_max_chars": 50000,
        "csv_preview_max_rows": 300,
    },
    "business": {
        "merge_pdf_max_files": 100,
        "merge_pdf_max_total_mb": 500,
        "merge_pdf_max_total_bytes": 500 * 1024 * 1024,
        "merge_pdf_max_total_pages": 5000,

        "pdf_general_max_mb": 300,
        "pdf_general_max_bytes": 300 * 1024 * 1024,
        "pdf_general_max_pages": 5000,

        "pdf_convert_max_mb": 250,
        "pdf_convert_max_bytes": 250 * 1024 * 1024,
        "pdf_convert_max_pages": 1500,

        "jpg_to_pdf_max_files": 200,
        "jpg_to_pdf_max_total_mb": 500,
        "jpg_to_pdf_max_total_bytes": 500 * 1024 * 1024,

        "heic_single_max_mb": 100,
        "heic_single_max_bytes": 100 * 1024 * 1024,
        "heic_bulk_max_files": 200,
        "heic_bulk_max_total_mb": 1000,
        "heic_bulk_max_total_bytes": 1000 * 1024 * 1024,

        "image_general_max_mb": 200,
        "image_general_max_bytes": 200 * 1024 * 1024,

        "office_general_max_mb": 250,
        "office_general_max_bytes": 250 * 1024 * 1024,

        "text_preview_max_chars": 100000,
        "csv_preview_max_rows": 1000,
    },
}

# --------------------------------------------------
# UNIT DATA
# --------------------------------------------------
UNIT_GROUPS = {
    "length": {
        "Millimeter (mm)": 0.001,
        "Centimeter (cm)": 0.01,
        "Meter (m)": 1.0,
        "Kilometer (km)": 1000.0,
        "Inch (in)": 0.0254,
        "Foot (ft)": 0.3048,
        "Yard (yd)": 0.9144,
        "Mile (mi)": 1609.344,
    },
    "area": {
        "Square Millimeter (mm²)": 1e-6,
        "Square Centimeter (cm²)": 1e-4,
        "Square Meter (m²)": 1.0,
        "Hectare (ha)": 10000.0,
        "Dekar": 1000.0,
        "Square Kilometer (km²)": 1e6,
        "Square Foot (ft²)": 0.09290304,
        "Acre": 4046.8564224,
    },
    "volume": {
        "Milliliter (mL)": 0.001,
        "Liter (L)": 1.0,
        "Cubic Meter (m³)": 1000.0,
        "Cubic Centimeter (cm³)": 0.001,
        "US Gallon": 3.785411784,
    },
    "weight": {
        "Milligram (mg)": 1e-6,
        "Gram (g)": 0.001,
        "Kilogram (kg)": 1.0,
        "Ton": 1000.0,
        "Pound (lb)": 0.45359237,
        "Ounce (oz)": 0.028349523125,
    },
    "speed": {
        "Meter/second (m/s)": 1.0,
        "Kilometer/hour (km/h)": 0.2777777778,
        "Mile/hour (mph)": 0.44704,
        "Knot": 0.5144444444,
    },
    "pressure": {
        "Pascal (Pa)": 1.0,
        "Kilopascal (kPa)": 1000.0,
        "Bar": 100000.0,
        "Atmosphere (atm)": 101325.0,
        "PSI": 6894.757293168,
    },
    "angle": {
        "Degree": 1.0,
        "Radian": 57.29577951308232,
        "Grad": 0.9,
    },
    "energy": {
        "Joule (J)": 1.0,
        "Kilojoule (kJ)": 1000.0,
        "Calorie (cal)": 4.184,
        "Kilocalorie (kcal)": 4184.0,
        "Watt-hour (Wh)": 3600.0,
        "Kilowatt-hour (kWh)": 3600000.0,
    },
    "time": {
        "Second": 1.0,
        "Minute": 60.0,
        "Hour": 3600.0,
        "Day": 86400.0,
        "Week": 604800.0,
    },
}
TEMPERATURE_UNITS = {"Celsius": None, "Fahrenheit": None, "Kelvin": None}

# --------------------------------------------------
# HELPERS
# --------------------------------------------------
def get_current_plan(request: Request) -> str:
    plan = (request.query_params.get("plan") or "free").strip().lower()
    if plan not in PLAN_LIMITS:
        plan = "free"
    return plan


def get_plan_limits(request: Request) -> dict:
    return PLAN_LIMITS[get_current_plan(request)]


def premium_upgrade_message(tool_name: str) -> str:
    return (
        f"Free plan limit reached for {tool_name}. "
        f"Upgrade to Pro for larger files, more pages, and batch processing."
    )


def get_units_for_category(category: str):
    if category == "temperature":
        return TEMPERATURE_UNITS
    return UNIT_GROUPS.get(category, UNIT_GROUPS["length"])


def convert_temperature(value, from_unit, to_unit):
    if from_unit == "Celsius":
        c = value
    elif from_unit == "Fahrenheit":
        c = (value - 32) * 5 / 9
    elif from_unit == "Kelvin":
        c = value - 273.15
    else:
        c = value

    if to_unit == "Celsius":
        return c
    if to_unit == "Fahrenheit":
        return (c * 9 / 5) + 32
    if to_unit == "Kelvin":
        return c + 273.15
    return c


def parse_page_ranges(pages_text: str, total_pages: int) -> list[int]:
    pages_text = (pages_text or "").replace(" ", "")
    if not pages_text:
        raise ValueError("Please enter page numbers.")

    selected = set()
    for part in pages_text.split(","):
        if not part:
            continue
        if "-" in part:
            start_str, end_str = part.split("-", 1)
            start = int(start_str)
            end = int(end_str)
            if start > end:
                start, end = end, start
            for p in range(start, end + 1):
                if 1 <= p <= total_pages:
                    selected.add(p - 1)
        else:
            p = int(part)
            if 1 <= p <= total_pages:
                selected.add(p - 1)

    if not selected:
        raise ValueError("No valid pages selected.")

    return sorted(selected)


def extract_youtube_video_id(url: str) -> str | None:
    try:
        parsed = urlparse(url.strip())
        host = parsed.netloc.lower()

        if "youtu.be" in host:
            vid = parsed.path.strip("/")
            return vid or None

        if "youtube.com" in host or "www.youtube.com" in host or "m.youtube.com" in host:
            if parsed.path == "/watch":
                qs = parse_qs(parsed.query)
                return qs.get("v", [None])[0]
            if parsed.path.startswith("/shorts/"):
                return parsed.path.split("/shorts/")[-1].split("/")[0]
            if parsed.path.startswith("/embed/"):
                return parsed.path.split("/embed/")[-1].split("/")[0]

        return None
    except Exception:
        return None


def clean_filename(name: str, fallback: str = "file") -> str:
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name or "")
    return name or fallback


def remove_file_safely(path: Path | None):
    if path:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass


def validate_size_by_plan(size_bytes: int, max_bytes: int, tool_name: str):
    if size_bytes > max_bytes:
        raise ValueError(premium_upgrade_message(tool_name))


def validate_pages_by_plan(page_count: int, max_pages: int, tool_name: str):
    if page_count > max_pages:
        raise ValueError(premium_upgrade_message(tool_name))


def validate_count_by_plan(item_count: int, max_count: int, tool_name: str):
    if item_count > max_count:
        raise ValueError(premium_upgrade_message(tool_name))


def validate_pdf_size(path: Path, max_bytes: int, tool_name: str = "PDF tool"):
    if path.stat().st_size > max_bytes:
        raise ValueError(premium_upgrade_message(tool_name))


def validate_pdf_pages(reader: PdfReader, max_pages: int, tool_name: str = "PDF tool"):
    if len(reader.pages) > max_pages:
        raise ValueError(premium_upgrade_message(tool_name))


# --------------------------------------------------
# HOME
# --------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "plan": get_current_plan(request),
        },
    )


@app.get("/ping")
async def ping():
    return {"status": "ok"}


@app.get("/pricing", response_class=HTMLResponse)
async def pricing_page(request: Request):
    return templates.TemplateResponse(
        "pricing.html",
        {
            "request": request,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


# --------------------------------------------------
# MAIN CATEGORY PAGES
# --------------------------------------------------
@app.get("/pdf", response_class=HTMLResponse)
async def pdf_page(request: Request):
    return templates.TemplateResponse(
        "pdf.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.get("/image", response_class=HTMLResponse)
async def image_page(request: Request):
    return templates.TemplateResponse(
        "image.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.get("/office", response_class=HTMLResponse)
async def office_page(request: Request):
    return templates.TemplateResponse(
        "office.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.get("/units", response_class=HTMLResponse)
async def units_page(request: Request):
    category = "length"
    all_units = {
        "temperature": list(TEMPERATURE_UNITS.keys()),
        **{k: list(v.keys()) for k, v in UNIT_GROUPS.items()},
    }
    return templates.TemplateResponse(
        "units.html",
        {
            "request": request,
            "result": None,
            "selected_category": category,
            "units": get_units_for_category(category),
            "all_units": all_units,
            "error": None,
            "plan": get_current_plan(request),
        },
    )


@app.post("/units", response_class=HTMLResponse)
async def convert_units(
    request: Request,
    category: str = Form(...),
    from_unit: str = Form(...),
    to_unit: str = Form(...),
    value: float = Form(...),
):
    try:
        units = get_units_for_category(category)
        if category == "temperature":
            result = convert_temperature(value, from_unit, to_unit)
        else:
            base_value = value * units[from_unit]
            result = base_value / units[to_unit]

        all_units = {
            "temperature": list(TEMPERATURE_UNITS.keys()),
            **{k: list(v.keys()) for k, v in UNIT_GROUPS.items()},
        }

        return templates.TemplateResponse(
            "units.html",
            {
                "request": request,
                "result": f"{result:.10f}".rstrip("0").rstrip("."),
                "selected_category": category,
                "units": units,
                "all_units": all_units,
                "error": None,
                "from_unit": from_unit,
                "to_unit": to_unit,
                "value": value,
                "plan": get_current_plan(request),
            },
        )
    except Exception as e:
        all_units = {
            "temperature": list(TEMPERATURE_UNITS.keys()),
            **{k: list(v.keys()) for k, v in UNIT_GROUPS.items()},
        }
        return templates.TemplateResponse(
            "units.html",
            {
                "request": request,
                "result": None,
                "selected_category": category,
                "units": get_units_for_category(category),
                "all_units": all_units,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
            },
            status_code=400,
        )


@app.get("/utility", response_class=HTMLResponse)
async def utility_page(request: Request):
    return templates.TemplateResponse(
        "utility.html",
        {
            "request": request,
            "result": None,
            "error": None,
            "plan": get_current_plan(request),
        },
    )


# --------------------------------------------------
# PDF TOOLS
# --------------------------------------------------
@app.get("/merge-pdf", response_class=HTMLResponse)
async def merge_pdf_page(request: Request):
    return templates.TemplateResponse(
        "tools/merge_pdf.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/merge-pdf", response_class=HTMLResponse)
@app.post("/merge", response_class=HTMLResponse)
async def merge_pdfs(request: Request, files: list[UploadFile] = File(...)):
    writer = PdfWriter()
    temp_files = []

    try:
        limits = get_plan_limits(request)

        valid_count = 0
        total_bytes = 0
        total_pages = 0

        for uploaded in files:
            if not uploaded.filename or not uploaded.filename.lower().endswith(".pdf"):
                continue

            temp_path = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(uploaded.filename)}"
            with open(temp_path, "wb") as buffer:
                shutil.copyfileobj(uploaded.file, buffer)

            if temp_path.stat().st_size == 0:
                remove_file_safely(temp_path)
                continue

            total_bytes += temp_path.stat().st_size
            if total_bytes > limits["merge_pdf_max_total_bytes"]:
                raise ValueError(premium_upgrade_message("Merge PDF"))

            temp_files.append(temp_path)
            valid_count += 1

            if valid_count > limits["merge_pdf_max_files"]:
                raise ValueError(premium_upgrade_message("Merge PDF"))

        if valid_count < 2:
            return templates.TemplateResponse(
                "tools/merge_pdf.html",
                {
                    "request": request,
                    "error": "Select at least 2 valid PDF files.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        for temp_path in temp_files:
            reader = PdfReader(str(temp_path))
            total_pages += len(reader.pages)

            if total_pages > limits["merge_pdf_max_total_pages"]:
                raise ValueError(premium_upgrade_message("Merge PDF"))

            for page in reader.pages:
                writer.add_page(page)

        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_buffer.seek(0)

        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="merged.pdf"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/merge_pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        for temp_path in temp_files:
            remove_file_safely(temp_path)


@app.get("/split-pdf", response_class=HTMLResponse)
async def split_pdf_page(request: Request):
    return templates.TemplateResponse(
        "tools/split_pdf.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/split-pdf", response_class=HTMLResponse)
@app.post("/split", response_class=HTMLResponse)
async def split_pdf(request: Request, file: UploadFile = File(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "tools/split_pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_general_max_bytes"], "Split PDF")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_general_max_pages"], "Split PDF")

        writer = PdfWriter()
        writer.add_page(reader.pages[0])

        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_buffer.seek(0)

        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="split_first_page.pdf"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/split_pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.get("/compress-pdf", response_class=HTMLResponse)
async def compress_pdf_page(request: Request):
    return templates.TemplateResponse(
        "tools/compress_pdf.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/compress-pdf", response_class=HTMLResponse)
@app.post("/pdf-compress", response_class=HTMLResponse)
async def pdf_compress(request: Request, file: UploadFile = File(...)):
    temp_input = None
    output_path = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "tools/compress_pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        output_path = TEMP_DIR / f"compressed_{uuid.uuid4().hex}.pdf"

        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_general_max_bytes"], "Compress PDF")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_general_max_pages"], "Compress PDF")

        doc = fitz.open(str(temp_input))
        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        return FileResponse(
            path=str(output_path),
            media_type="application/pdf",
            filename="compressed.pdf",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/compress_pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/jpg-to-pdf", response_class=HTMLResponse)
async def jpg_to_pdf(request: Request, files: list[UploadFile] = File(...)):
    try:
        limits = get_plan_limits(request)
        validate_count_by_plan(len(files), limits["jpg_to_pdf_max_files"], "JPG to PDF")

        images = []
        total_bytes = 0

        for uploaded in files:
            data = await uploaded.read()
            if data:
                total_bytes += len(data)
                if total_bytes > limits["jpg_to_pdf_max_total_bytes"]:
                    raise ValueError(premium_upgrade_message("JPG to PDF"))

                img = Image.open(io.BytesIO(data)).convert("RGB")
                images.append(img)

        if not images:
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Please select valid images.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        pdf_buffer = io.BytesIO()
        first, rest = images[0], images[1:]
        first.save(pdf_buffer, format="PDF", save_all=True, append_images=rest)
        pdf_buffer.seek(0)

        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="images_to_pdf.pdf"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.get("/pdf-to-word", response_class=HTMLResponse)
async def pdf_to_word_page(request: Request):
    return templates.TemplateResponse(
        "tools/pdf_to_word.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/pdf-to-word", response_class=HTMLResponse)
async def pdf_to_word(request: Request, file: UploadFile = File(...)):
    temp_input = None
    output_path = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "tools/pdf_to_word.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_convert_max_bytes"], "PDF to Word")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_convert_max_pages"], "PDF to Word")

        doc = Document()
        doc.add_heading("PDF to Word Output", level=1)

        pdf_doc = fitz.open(str(temp_input))
        for page_no, page in enumerate(pdf_doc, start=1):
            text = page.get_text("text") or ""
            doc.add_heading(f"Page {page_no}", level=2)

            if text.strip():
                for line in text.splitlines():
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph("[No extractable text found on this page]")

        pdf_doc.close()

        output_path = TEMP_DIR / f"pdf_to_word_{uuid.uuid4().hex}.docx"
        doc.save(str(output_path))

        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="pdf_to_word.docx",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/pdf_to_word.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/pdf-to-excel", response_class=HTMLResponse)
async def pdf_to_excel(request: Request, file: UploadFile = File(...)):
    temp_input = None
    output_path = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_convert_max_bytes"], "PDF to Excel")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_convert_max_pages"], "PDF to Excel")

        wb = openpyxl.Workbook()
        default_sheet = wb.active
        wb.remove(default_sheet)

        found_table = False
        with pdfplumber.open(str(temp_input)) as pdf:
            for page_no, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables, start=1):
                    found_table = True
                    ws = wb.create_sheet(title=f"P{page_no}_T{table_idx}")
                    for row in table:
                        ws.append([cell if cell is not None else "" for cell in row])

        if not found_table:
            ws = wb.create_sheet(title="NoTable")
            ws.append(["No extractable table found in this PDF."])

        output_path = TEMP_DIR / f"pdf_to_excel_{uuid.uuid4().hex}.xlsx"
        wb.save(str(output_path))

        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename="pdf_to_excel.xlsx",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/pdf-to-jpg", response_class=HTMLResponse)
async def pdf_to_jpg(request: Request, file: UploadFile = File(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_convert_max_bytes"], "PDF to JPG")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_convert_max_pages"], "PDF to JPG")

        pdf_doc = fitz.open(str(temp_input))
        page = pdf_doc.load_page(0)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_bytes = pix.tobytes("jpg")
        pdf_doc.close()

        return Response(
            content=img_bytes,
            media_type="image/jpeg",
            headers={"Content-Disposition": 'attachment; filename="pdf_first_page.jpg"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/pdf-rotate", response_class=HTMLResponse)
async def pdf_rotate(request: Request, file: UploadFile = File(...), angle: int = Form(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        if angle not in [90, 180, 270]:
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Angle must be 90, 180 or 270.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_general_max_bytes"], "Rotate PDF")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_general_max_pages"], "Rotate PDF")

        writer = PdfWriter()
        for page in reader.pages:
            page.rotate(angle)
            writer.add_page(page)

        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_buffer.seek(0)

        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="rotated.pdf"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/pdf-remove-pages", response_class=HTMLResponse)
async def pdf_remove_pages(request: Request, file: UploadFile = File(...), pages: str = Form(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".pdf"):
            return templates.TemplateResponse(
                "pdf.html",
                {
                    "request": request,
                    "error": "Please upload a PDF file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_pdf_size(temp_input, limits["pdf_general_max_bytes"], "Remove PDF Pages")
        reader = PdfReader(str(temp_input))
        validate_pdf_pages(reader, limits["pdf_general_max_pages"], "Remove PDF Pages")

        total_pages = len(reader.pages)
        pages_to_remove = set(parse_page_ranges(pages, total_pages))
        remaining_pages = total_pages - len(pages_to_remove)
        if remaining_pages <= 0:
            raise ValueError("You cannot remove all pages from the PDF.")

        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if i not in pages_to_remove:
                writer.add_page(page)

        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_buffer.seek(0)

        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="removed_pages.pdf"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


# --------------------------------------------------
# IMAGE TOOLS
# --------------------------------------------------
@app.get("/heic-to-jpg", response_class=HTMLResponse)
async def heic_to_jpg_page(request: Request):
    return templates.TemplateResponse(
        "tools/heic_to_jpg.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/heic-to-jpg", response_class=HTMLResponse)
async def heic_to_jpg(request: Request, file: UploadFile = File(...)):
    try:
        limits = get_plan_limits(request)

        data = await file.read()
        validate_size_by_plan(len(data), limits["heic_single_max_bytes"], "HEIC to JPG")

        img = Image.open(io.BytesIO(data)).convert("RGB")
        output = io.BytesIO()
        img.save(output, "JPEG", quality=95)
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/jpeg",
            headers={"Content-Disposition": 'attachment; filename="converted.jpg"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/heic_to_jpg.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.get("/bulk-heic-to-jpg", response_class=HTMLResponse)
async def bulk_heic_to_jpg_page(request: Request):
    limits = get_plan_limits(request)
    return templates.TemplateResponse(
        "tools/bulk_heic_to_jpg.html",
        {
            "request": request,
            "error": None,
            "max_files": limits["heic_bulk_max_files"],
            "max_total_mb": limits["heic_bulk_max_total_mb"],
            "plan": get_current_plan(request),
            "limits": limits,
        },
    )


@app.post("/bulk-heic-to-jpg", response_class=HTMLResponse)
async def bulk_heic_to_jpg(request: Request, files: list[UploadFile] = File(...)):
    try:
        limits = get_plan_limits(request)
        valid_files = [f for f in files if f.filename]

        if not valid_files:
            return templates.TemplateResponse(
                "tools/bulk_heic_to_jpg.html",
                {
                    "request": request,
                    "error": "Please select at least 1 HEIC file.",
                    "max_files": limits["heic_bulk_max_files"],
                    "max_total_mb": limits["heic_bulk_max_total_mb"],
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        if len(valid_files) > limits["heic_bulk_max_files"]:
            return templates.TemplateResponse(
                "tools/bulk_heic_to_jpg.html",
                {
                    "request": request,
                    "error": premium_upgrade_message("Bulk HEIC to JPG"),
                    "max_files": limits["heic_bulk_max_files"],
                    "max_total_mb": limits["heic_bulk_max_total_mb"],
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        file_items = []
        total_bytes = 0

        for uploaded in valid_files:
            data = await uploaded.read()
            if not data:
                continue

            total_bytes += len(data)
            if total_bytes > limits["heic_bulk_max_total_bytes"]:
                return templates.TemplateResponse(
                    "tools/bulk_heic_to_jpg.html",
                    {
                        "request": request,
                        "error": premium_upgrade_message("Bulk HEIC to JPG"),
                        "max_files": limits["heic_bulk_max_files"],
                        "max_total_mb": limits["heic_bulk_max_total_mb"],
                        "plan": get_current_plan(request),
                        "limits": limits,
                    },
                    status_code=400,
                )

            filename = clean_filename(Path(uploaded.filename).stem, "image")
            file_items.append((filename, data))

        if not file_items:
            return templates.TemplateResponse(
                "tools/bulk_heic_to_jpg.html",
                {
                    "request": request,
                    "error": "No valid files found.",
                    "max_files": limits["heic_bulk_max_files"],
                    "max_total_mb": limits["heic_bulk_max_total_mb"],
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
            for filename, data in file_items:
                img = Image.open(io.BytesIO(data)).convert("RGB")
                jpg_buffer = io.BytesIO()
                img.save(jpg_buffer, format="JPEG", quality=95)
                jpg_buffer.seek(0)
                zip_file.writestr(f"{filename}.jpg", jpg_buffer.getvalue())

        zip_buffer.seek(0)

        return Response(
            content=zip_buffer.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition": 'attachment; filename="bulk_heic_to_jpg.zip"'},
        )

    except Exception as e:
        limits = get_plan_limits(request)
        return templates.TemplateResponse(
            "tools/bulk_heic_to_jpg.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "max_files": limits["heic_bulk_max_files"],
                "max_total_mb": limits["heic_bulk_max_total_mb"],
                "plan": get_current_plan(request),
                "limits": limits,
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.get("/image-compressor", response_class=HTMLResponse)
async def image_compressor_page(request: Request):
    return templates.TemplateResponse(
        "tools/image_compressor.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/image-compressor", response_class=HTMLResponse)
@app.post("/compress-image", response_class=HTMLResponse)
async def compress_image(request: Request, file: UploadFile = File(...), quality: int = Form(70)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Image Compressor")

        img = Image.open(io.BytesIO(data)).convert("RGB")
        output = io.BytesIO()
        img.save(output, "JPEG", quality=max(10, min(95, quality)), optimize=True)
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/jpeg",
            headers={"Content-Disposition": 'attachment; filename="compressed.jpg"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/image_compressor.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/convert-image", response_class=HTMLResponse)
async def convert_image(request: Request, file: UploadFile = File(...), target_format: str = Form(...)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Image Converter")

        img = Image.open(io.BytesIO(data))
        fmt = target_format.lower()
        output = io.BytesIO()

        if fmt in ["jpg", "jpeg"]:
            img = img.convert("RGB")
            img.save(output, "JPEG", quality=95)
            media_type = "image/jpeg"
            filename = "converted.jpg"
        elif fmt == "png":
            img = img.convert("RGBA")
            img.save(output, "PNG")
            media_type = "image/png"
            filename = "converted.png"
        elif fmt == "webp":
            img.save(output, "WEBP", quality=95)
            media_type = "image/webp"
            filename = "converted.webp"
        else:
            return templates.TemplateResponse(
                "image.html",
                {
                    "request": request,
                    "error": "Unsupported target format.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        output.seek(0)
        return Response(
            content=output.getvalue(),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "image.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/resize-image", response_class=HTMLResponse)
async def resize_image(
    request: Request,
    file: UploadFile = File(...),
    width: int = Form(...),
    height: int = Form(...),
):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Resize Image")

        img = Image.open(io.BytesIO(data)).convert("RGB")
        resized = img.resize((width, height))
        output = io.BytesIO()
        resized.save(output, "JPEG", quality=95)
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/jpeg",
            headers={"Content-Disposition": 'attachment; filename="resized.jpg"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "image.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/remove-bg-basic", response_class=HTMLResponse)
async def remove_bg_basic(request: Request, file: UploadFile = File(...), threshold: int = Form(240)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Background Remover")

        img = Image.open(io.BytesIO(data)).convert("RGBA")
        pixels = img.getdata()
        new_pixels = []
        threshold = max(0, min(255, threshold))

        for r, g, b, a in pixels:
            if r >= threshold and g >= threshold and b >= threshold:
                new_pixels.append((255, 255, 255, 0))
            else:
                new_pixels.append((r, g, b, a))

        img.putdata(new_pixels)
        output = io.BytesIO()
        img.save(output, "PNG")
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/png",
            headers={"Content-Disposition": 'attachment; filename="background_removed.png"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "image.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/image-to-svg", response_class=HTMLResponse)
async def image_to_svg(request: Request, file: UploadFile = File(...)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Image to SVG")

        img = Image.open(io.BytesIO(data))
        width, height = img.size
        mime = file.content_type or "image/png"
        b64 = base64.b64encode(data).decode("utf-8")

        svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
  <image href="data:{mime};base64,{b64}" x="0" y="0" width="{width}" height="{height}" />
</svg>"""

        return Response(
            content=svg.encode("utf-8"),
            media_type="image/svg+xml",
            headers={"Content-Disposition": 'attachment; filename="image.svg"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "image.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/reduce-image-size", response_class=HTMLResponse)
async def reduce_image_size(request: Request, file: UploadFile = File(...), quality: int = Form(50)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Reduce Image Size")

        img = Image.open(io.BytesIO(data)).convert("RGB")
        output = io.BytesIO()
        img.save(output, "WEBP", quality=max(10, min(95, quality)), method=6)
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/webp",
            headers={"Content-Disposition": 'attachment; filename="reduced.webp"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "image.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


# --------------------------------------------------
# OFFICE TOOLS
# --------------------------------------------------
@app.post("/csv-to-excel", response_class=HTMLResponse)
async def csv_to_excel(request: Request, file: UploadFile = File(...)):
    temp_input = None
    output_path = None

    try:
        limits = get_plan_limits(request)

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_size_by_plan(temp_input.stat().st_size, limits["office_general_max_bytes"], "CSV to Excel")

        df = pd.read_csv(temp_input)
        output_path = TEMP_DIR / f"{uuid.uuid4().hex}.xlsx"
        df.to_excel(output_path, index=False)

        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename="converted.xlsx",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "office.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/excel-to-csv", response_class=HTMLResponse)
async def excel_to_csv(request: Request, file: UploadFile = File(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_size_by_plan(temp_input.stat().st_size, limits["office_general_max_bytes"], "Excel to CSV")

        df = pd.read_excel(temp_input)
        csv_text = df.to_csv(index=False)

        return Response(
            content=csv_text.encode("utf-8"),
            media_type="text/csv",
            headers={"Content-Disposition": 'attachment; filename="converted.csv"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "office.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.post("/txt-to-docx", response_class=HTMLResponse)
async def txt_to_docx(request: Request, file: UploadFile = File(...)):
    output_path = None

    try:
        limits = get_plan_limits(request)

        data = await file.read()
        validate_size_by_plan(len(data), limits["office_general_max_bytes"], "TXT to DOCX")

        text = data.decode("utf-8", errors="ignore")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)

        output_path = TEMP_DIR / f"{uuid.uuid4().hex}.docx"
        doc.save(str(output_path))

        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="converted.docx",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "office.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )


@app.post("/docx-to-txt", response_class=HTMLResponse)
async def docx_to_txt(request: Request, file: UploadFile = File(...)):
    temp_input = None

    try:
        limits = get_plan_limits(request)

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_size_by_plan(temp_input.stat().st_size, limits["office_general_max_bytes"], "DOCX to TXT")

        doc = Document(str(temp_input))
        text = "\n".join([p.text for p in doc.paragraphs])

        return Response(
            content=text.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": 'attachment; filename="converted.txt"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "office.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


@app.get("/word-to-pdf", response_class=HTMLResponse)
async def word_to_pdf_page(request: Request):
    return templates.TemplateResponse(
        "tools/word_to_pdf.html",
        {
            "request": request,
            "error": None,
            "plan": get_current_plan(request),
            "limits": get_plan_limits(request),
        },
    )


@app.post("/word-to-pdf", response_class=HTMLResponse)
@app.post("/docx-to-pdf", response_class=HTMLResponse)
async def docx_to_pdf(request: Request, file: UploadFile = File(...)):
    temp_input = None
    output_path = None

    try:
        limits = get_plan_limits(request)

        if not file.filename or not file.filename.lower().endswith(".docx"):
            return templates.TemplateResponse(
                "tools/word_to_pdf.html",
                {
                    "request": request,
                    "error": "Please upload a DOCX file.",
                    "plan": get_current_plan(request),
                    "limits": limits,
                },
                status_code=400,
            )

        temp_input = TEMP_DIR / f"{uuid.uuid4().hex}_{clean_filename(file.filename)}"
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        validate_size_by_plan(temp_input.stat().st_size, limits["office_general_max_bytes"], "Word to PDF")

        doc = Document(str(temp_input))
        output_path = TEMP_DIR / f"{uuid.uuid4().hex}.pdf"

        c = canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4
        left_margin = 50
        top_margin = height - 50
        bottom_margin = 50
        y = top_margin

        font_name = "Helvetica"
        if DEJAVU_FONT_PATH.exists():
            font_name = "DejaVuSans"

        c.setFont(font_name, 11)
        max_chars = 95

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if not text:
                y -= 16
                if y < bottom_margin:
                    c.showPage()
                    c.setFont(font_name, 11)
                    y = top_margin
                continue

            chunks = [text[i:i + max_chars] for i in range(0, len(text), max_chars)]

            for chunk in chunks:
                c.drawString(left_margin, y, chunk)
                y -= 18

                if y < bottom_margin:
                    c.showPage()
                    c.setFont(font_name, 11)
                    y = top_margin

        c.save()

        return FileResponse(
            path=str(output_path),
            media_type="application/pdf",
            filename="converted.pdf",
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/word_to_pdf.html",
            {
                "request": request,
                "error": f"Error: {str(e)}",
                "plan": get_current_plan(request),
                "limits": get_plan_limits(request),
            },
            status_code=500 if not isinstance(e, ValueError) else 400,
        )
    finally:
        remove_file_safely(temp_input)


# --------------------------------------------------
# UTILITY TOOLS
# --------------------------------------------------
@app.get("/uuid-generator", response_class=HTMLResponse)
async def uuid_generator_page(request: Request):
    return templates.TemplateResponse(
        "tools/uuid_generator.html",
        {"request": request, "error": None, "plan": get_current_plan(request)},
    )


@app.get("/qr-code-generator", response_class=HTMLResponse)
async def qr_generator_page(request: Request):
    return templates.TemplateResponse(
        "tools/qr_generator.html",
        {"request": request, "error": None, "plan": get_current_plan(request)},
    )


@app.get("/password-generator", response_class=HTMLResponse)
async def password_generator_page(request: Request):
    return templates.TemplateResponse(
        "tools/password_generator.html",
        {"request": request, "error": None, "result": None, "plan": get_current_plan(request)},
    )


@app.post("/password-generator", response_class=HTMLResponse)
@app.post("/password-generate", response_class=HTMLResponse)
async def password_generate(request: Request, length: int = Form(16), use_symbols: str = Form("true")):
    try:
        length = max(6, min(128, int(length)))
        chars = string.ascii_letters + string.digits
        if str(use_symbols).lower() == "true":
            chars += "!@#$%^&*()-_=+[]{};:,.?/"

        pwd = "".join(secrets.choice(chars) for _ in range(length))

        return templates.TemplateResponse(
            "tools/password_generator.html",
            {"request": request, "result": pwd, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/password_generator.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
            status_code=500,
        )


@app.get("/json-formatter", response_class=HTMLResponse)
async def json_formatter_page(request: Request):
    return templates.TemplateResponse(
        "tools/json_formatter.html",
        {"request": request, "error": None, "result": None, "plan": get_current_plan(request)},
    )


@app.post("/json-formatter", response_class=HTMLResponse)
async def json_formatter(request: Request, text: str = Form(...)):
    try:
        parsed = json.loads(text)
        result = json.dumps(parsed, indent=4, ensure_ascii=False)
        return templates.TemplateResponse(
            "tools/json_formatter.html",
            {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "tools/json_formatter.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/qr-generate", response_class=HTMLResponse)
async def qr_generate_utility(request: Request, text: str = Form(...)):
    try:
        img = qrcode.make(text)
        output = io.BytesIO()
        img.save(output, format="PNG")
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/png",
            headers={"Content-Disposition": 'attachment; filename="qr.png"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "error": f"Error: {str(e)}", "result": None, "plan": get_current_plan(request)},
            status_code=500,
        )


@app.post("/word-counter", response_class=HTMLResponse)
async def word_counter(request: Request, text: str = Form(...)):
    words = len(text.split())
    chars = len(text)
    lines = len(text.splitlines()) if text else 0
    reading_time = max(1, round(words / 200)) if words else 0
    result = f"Words: {words} | Characters: {chars} | Lines: {lines} | Reading time: {reading_time} min"

    return templates.TemplateResponse(
        "utility.html",
        {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
    )


@app.post("/base64-encode", response_class=HTMLResponse)
async def base64_encode_tool(request: Request, text: str = Form(...)):
    result = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    return templates.TemplateResponse(
        "utility.html",
        {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
    )


@app.post("/base64-decode", response_class=HTMLResponse)
async def base64_decode_tool(request: Request, text: str = Form(...)):
    try:
        result = base64.b64decode(text.encode("utf-8")).decode("utf-8", errors="ignore")
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/markdown-to-html", response_class=HTMLResponse)
async def markdown_to_html(request: Request, text: str = Form(...)):
    try:
        html = md.markdown(text, extensions=["tables", "fenced_code"])
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": html, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/youtube-thumbnail", response_class=HTMLResponse)
async def youtube_thumbnail(request: Request, url: str = Form(...)):
    try:
        video_id = extract_youtube_video_id(url)
        if not video_id:
            return templates.TemplateResponse(
                "utility.html",
                {"request": request, "result": None, "error": "Invalid YouTube URL.", "plan": get_current_plan(request)},
                status_code=400,
            )

        thumb_url = f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg"
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": thumb_url, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/case-converter", response_class=HTMLResponse)
async def case_converter(request: Request, text: str = Form(...), mode: str = Form(...)):
    try:
        if mode == "upper":
            result = text.upper()
        elif mode == "lower":
            result = text.lower()
        elif mode == "title":
            result = text.title()
        elif mode == "snake":
            result = re.sub(r"\W+", "_", text.strip().lower()).strip("_")
        elif mode == "camel":
            parts = re.sub(r"\W+", " ", text).split()
            result = parts[0].lower() + "".join(word.capitalize() for word in parts[1:]) if parts else ""
        else:
            result = text

        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/age-calculator", response_class=HTMLResponse)
async def age_calculator(request: Request, birthdate: str = Form(...)):
    try:
        born = datetime.strptime(birthdate, "%Y-%m-%d").date()
        today = date.today()
        years = today.year - born.year - ((today.month, today.day) < (born.month, born.day))
        days = (today - born).days
        result = f"Age: {years} years | Total days: {days}"

        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": result, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/favicon-generator", response_class=HTMLResponse)
async def favicon_generator(request: Request, file: UploadFile = File(...)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["image_general_max_bytes"], "Favicon Generator")

        img = Image.open(io.BytesIO(data)).convert("RGBA")
        img = img.resize((64, 64))
        output = io.BytesIO()
        img.save(output, format="ICO", sizes=[(16, 16), (32, 32), (48, 48), (64, 64)])
        output.seek(0)

        return Response(
            content=output.getvalue(),
            media_type="image/x-icon",
            headers={"Content-Disposition": 'attachment; filename="favicon.ico"'},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/csv-viewer", response_class=HTMLResponse)
async def csv_viewer(request: Request, file: UploadFile = File(...)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["office_general_max_bytes"], "CSV Viewer")

        df = pd.read_csv(io.BytesIO(data))
        html_table = df.head(limits["csv_preview_max_rows"]).to_html(classes="table-view", index=False, border=0)

        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": html_table, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


@app.post("/kml-viewer", response_class=HTMLResponse)
async def kml_viewer(request: Request, file: UploadFile = File(...)):
    try:
        limits = get_plan_limits(request)
        data = await file.read()
        validate_size_by_plan(len(data), limits["office_general_max_bytes"], "KML Viewer")

        text = data.decode("utf-8", errors="ignore")
        preview = text[:limits["text_preview_max_chars"]]

        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": preview, "error": None, "plan": get_current_plan(request)},
        )
    except Exception as e:
        return templates.TemplateResponse(
            "utility.html",
            {"request": request, "result": None, "error": f"Error: {str(e)}", "plan": get_current_plan(request)},
        )


# --------------------------------------------------
# SEO PAGES
# --------------------------------------------------
@app.get("/free-pdf-tools", response_class=HTMLResponse)
async def free_pdf_tools(request: Request):
    return templates.TemplateResponse(
        "seo/seo_page.html",
        {
            "request": request,
            "title": "Free PDF Tools Online",
            "heading": "Free PDF Tools Online",
            "subtitle": "Merge, split, compress and convert PDF files easily.",
            "description": "Use free PDF tools online including merge, split, compress and convert PDF files quickly.",
            "plan": get_current_plan(request),
        },
    )


@app.get("/free-image-tools", response_class=HTMLResponse)
async def free_image_tools(request: Request):
    return templates.TemplateResponse(
        "seo/seo_page.html",
        {
            "request": request,
            "title": "Free Image Tools Online",
            "heading": "Free Image Tools",
            "subtitle": "Compress, convert and resize images online.",
            "description": "Free image tools including image compressor, HEIC converter and resize tools.",
            "plan": get_current_plan(request),
        },
    )


@app.get("/online-utility-tools", response_class=HTMLResponse)
async def utility_tools(request: Request):
    return templates.TemplateResponse(
        "seo/seo_page.html",
        {
            "request": request,
            "title": "Online Utility Tools",
            "heading": "Online Utility Tools",
            "subtitle": "Password generator, JSON formatter and more.",
            "description": "Use useful online utility tools including password generator and JSON formatter.",
            "plan": get_current_plan(request),
        },
    )


# --------------------------------------------------
# STATIC SEO FILES
# --------------------------------------------------
@app.get("/sitemap.xml", include_in_schema=False)
def sitemap():
    if not SITEMAP_PATH.exists():
        return Response(content="sitemap.xml not found", media_type="text/plain", status_code=404)
    return FileResponse(str(SITEMAP_PATH), media_type="application/xml")


@app.get("/robots.txt", include_in_schema=False)
def robots():
    if not ROBOTS_PATH.exists():
        return Response(content="robots.txt not found", media_type="text/plain", status_code=404)
    return FileResponse(str(ROBOTS_PATH), media_type="text/plain")


# --------------------------------------------------
# PROGRAMMATIC SEO UNIT PAGES
# --------------------------------------------------
SEO_PAGES_BY_SLUG = {page["slug"]: page for page in SEO_UNIT_PAGES}


@app.get("/convert/{slug}", response_class=HTMLResponse)
async def seo_unit_page(request: Request, slug: str):
    page = SEO_PAGES_BY_SLUG.get(slug)
    if not page:
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "plan": get_current_plan(request),
            },
            status_code=404,
        )

    return templates.TemplateResponse(
        "seo/unit_converter_landing.html",
        {
            "request": request,
            "page": page,
            "plan": get_current_plan(request),
        },
    )